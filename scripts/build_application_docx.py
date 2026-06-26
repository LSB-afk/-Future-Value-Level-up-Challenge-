from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "MoveValue_참가신청서_보강본.docx"
SCREENSHOT_DIR = ROOT / "deliverables" / "screenshots"


TOKENS = {
    "font": "Malgun Gothic",
    "body_size": 9.5,
    "heading_green": "0F766E",
    "heading_dark": "17212B",
    "border": "D8DDE2",
    "header_fill": "EAF7F0",
    "muted_fill": "F7F9FB",
    "warning_fill": "FFF7ED",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=95, bottom=70, end=95):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=None):
    color = color or TOKENS["border"]
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, size=None, bold=False, color=None):
    run.font.name = TOKENS["font"]
    run._element.rPr.rFonts.set(qn("w:eastAsia"), TOKENS["font"])
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para(paragraph, size=None):
    for run in paragraph.runs:
        set_run_font(run, size or TOKENS["body_size"])


def add_paragraph(doc, text="", bold=False, color=None, align=None):
    paragraph = doc.add_paragraph()
    if align:
        paragraph.alignment = align
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, TOKENS["body_size"], bold=bold, color=color)
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    size = 15 if level == 1 else 12 if level == 2 else 10.5
    color = TOKENS["heading_green"] if level <= 2 else TOKENS["heading_dark"]
    set_run_font(run, size, bold=True, color=color)
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_table(doc, headers, rows, widths=None, header_fill=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        set_cell_shading(cell, header_fill or TOKENS["header_fill"])
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(str(header))
        set_run_font(run, 8.8, bold=True, color=TOKENS["heading_dark"])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            run = cell.paragraphs[0].add_run(str(value))
            set_run_font(run, 8.6)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_kv_table(doc, rows):
    return add_table(doc, ["항목", "작성 내용"], rows, widths=[1.55, 5.05])


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run, TOKENS["body_size"])


def add_image(doc, filename, title, description, width=6.25):
    path = SCREENSHOT_DIR / filename
    if not path.exists():
        add_paragraph(doc, f"[캡처 누락] {filename}", bold=True, color="B45309")
        return
    add_heading(doc, title, 3)
    add_paragraph(doc, description)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = caption.add_run(f"그림. {title} ({filename})")
    set_run_font(cap_run, 8.2, color="64748B")


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    styles["Normal"].font.name = TOKENS["font"]
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), TOKENS["font"])
    styles["Normal"].font.size = Pt(TOKENS["body_size"])
    styles["Normal"].paragraph_format.line_spacing = 1.08
    styles["Normal"].paragraph_format.space_after = Pt(4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("2026 국토교통 서비스 발굴 경연 참가신청서")
    set_run_font(run, 18, bold=True, color=TOKENS["heading_dark"])
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("MoveValue: 주거-이동 통합 생활권 매칭 서비스")
    set_run_font(run, 12, bold=True, color=TOKENS["heading_green"])
    add_paragraph(
        doc,
        "주의: 연락처, 서명, 개인정보 수집·이용 동의는 제출 전 신청자가 직접 확인하고 기재해야 합니다.",
        bold=True,
        color="B45309",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_heading(doc, "1. 참가 신청서", 1)
    add_kv_table(
        doc,
        [
            ("신청 기관명", "이승보 / MoveValue 프로젝트팀(개인 참가, 예비창업)"),
            ("사업 분야", "국토교통 데이터 기반 생활권 추천·부동산·교통 연계 서비스"),
            ("신청 담당자명", "이승보"),
            ("담당자 연락처", "제출 전 기재"),
            ("담당자 이메일", "knightpen@icloud.com"),
            ("서비스 모델명", "MoveValue: 주거-이동 통합 생활권 매칭 서비스"),
            ("서비스 사업 분야", "부동산·교통·MaaS·생활 SOC 데이터 융합 서비스"),
            (
                "산출물(프로토타입)",
                "웹 기반 지도 대시보드, 아파트 매칭 API, 단지 상세 대시보드, 전세 위험 신호 점검, AI Agent 질의응답, 데이터/API 연계 문서, 화면 캡처 및 아키텍처 다이어그램",
            ),
            ("소요 예산", "20,000천원"),
            ("소요 기간", "2026.06.20 ~ 2026.08.31(약 2.5개월)"),
            ("주관 기업(역할)", "MoveValue 프로젝트팀: 서비스 기획, 프로토타입 개발, 데이터 결합·점수화"),
            ("참여/협력 기업", "협력 후보: 공공데이터 제공기관, 국토교통 데이터 오픈마켓, 부동산 플랫폼, MaaS/교통 데이터 사업자"),
        ],
    )

    add_heading(doc, "2. 국토교통 분야 실증서비스 모델 구현 기획서", 1)
    add_heading(doc, "서비스 모델명 및 분야", 2)
    add_paragraph(doc, "MoveValue: 주거-이동 통합 생활권 매칭 서비스")
    add_paragraph(doc, "분야: 국토교통 데이터 융합, 생활권 추천, 부동산·교통 연계 서비스")

    add_heading(doc, "서비스 제공 대상", 2)
    add_paragraph(
        doc,
        "공공과 민간 모두 해당합니다. 공공은 지자체 생활권 정책 부서와 국토교통 데이터 활용 사업을 대상으로 하고, 민간은 부동산 플랫폼, MaaS 플랫폼, 기업 이주 지원 서비스, 청년·신혼부부·직장인 사용자를 대상으로 합니다.",
    )

    add_heading(doc, "실증서비스 아이디어 모델 개요", 2)
    add_paragraph(
        doc,
        "MoveValue는 사용자의 월 주거 예산, 출퇴근 목적지, 생활 인프라 선호도, 가구 유형, 계약 조건을 기반으로 적합한 생활권과 주거 후보지를 추천합니다. 단순 부동산 가격 검색을 넘어 주거비, 이동시간, 교통 접근성, 생활 SOC, 안전·환경 지표, 전세 위험 신호를 통합 분석합니다.",
    )
    add_bullets(
        doc,
        [
            "사용자 조건 기반 생활권 및 아파트 단지 추천",
            "지도 기반 후보 지역 탐색과 아파트 단지 상세 대시보드",
            "실거래가·공시가격·전세가율·거래 추이 확인",
            "후보 지역/단지 비교와 전세 위험 신호 점검",
            "AI Agent 기반 설명형 추천과 계약 전 확인 항목 안내",
            "공공데이터/API 기반 데이터 연계 구조와 모바일 반응형 UI",
        ],
    )

    add_heading(doc, "현재 구현된 프로토타입 화면", 2)
    add_image(doc, "01-main-screen.png", "메인 지도 워크스페이스", "조건 입력과 서울 지도 기반 아파트 단지 탐색을 한 화면에 배치했습니다.")
    add_image(doc, "03-recommendation-results.png", "생활권/단지 추천 결과", "목적지와 예산을 입력하면 단지별 추천 순위, 월세, 통근시간을 보여줍니다.")
    add_image(doc, "05-property-dashboard.png", "아파트 상세 대시보드", "선택 단지의 기본정보, 가격, 거래 추이, 전세 위험 신호를 한 화면에서 검토합니다.")
    add_image(doc, "07-candidate-comparison.png", "후보 단지 비교", "즐겨찾기한 단지의 가격, 전세가율, 위험도, 생활 SOC를 표로 비교합니다.")
    add_image(doc, "09-ai-agent-qa.png", "AI Agent 질의응답", "질문에 대해 가격 근거, 통근 근거, 위험 근거, 주변 입지 근거를 분리해 답변합니다.")
    add_image(doc, "10-mobile-responsive.png", "모바일 반응형 화면", "모바일 폭에서도 조건 입력과 추천 카드가 1열로 유지됩니다.", width=2.35)

    add_heading(doc, "배경 및 필요성", 2)
    add_paragraph(
        doc,
        "주거 선택 과정에서 시민은 가격, 직장 접근성, 대중교통, 생활시설, 전세 위험 요소를 각각 다른 서비스에서 확인해야 합니다. 이로 인해 월세는 낮지만 통근시간과 교통비가 높아지는 선택, 생활 인프라 접근성이 떨어지는 선택, 계약 전 위험 신호를 놓치는 문제가 발생합니다.",
    )
    add_paragraph(
        doc,
        "국토교통 분야에는 실거래가, 공시가격, 건축물대장, 지오코딩, 도시철도, 생활 SOC 데이터가 존재합니다. MoveValue는 이 데이터를 사용자가 실제로 이해할 수 있는 지도, 점수, 설명형 Agent로 연결해 국민 체감형 주거 의사결정 서비스를 만드는 것을 목표로 합니다.",
    )

    add_heading(doc, "기존 서비스와 MoveValue 비교표", 2)
    add_table(
        doc,
        ["구분", "기존 부동산 서비스", "MoveValue"],
        [
            ("탐색 기준", "매물 가격, 위치, 면적 중심", "주거비, 통근시간, 생활 SOC, 안전환경, 전세 위험 신호 통합"),
            ("추천 방식", "사용자가 조건을 직접 필터링", "사용자 예산·목적지·가구 유형·가중치 기반 점수화"),
            ("설명 가능성", "가격과 매물 정보 중심", "가격/통근/위험/입지/확인 서류 근거 그룹 제공"),
            ("지도 UX", "매물 검색 지도", "생활권 추천과 단지 위험 점검을 결합한 의사결정 지도"),
            ("공공 활용", "민간 매물 검색 중심", "지자체 정책 리포트, B2B/B2G API 확장 가능"),
        ],
        widths=[1.25, 2.5, 2.85],
    )

    add_heading(doc, "사용자 문제점과 해결 기능 매핑표", 2)
    add_table(
        doc,
        ["사용자 문제", "MoveValue 해결 기능", "제공 가치"],
        [
            ("주거비와 통근 부담을 따로 계산해야 함", "예산·목적지 기반 추천 점수", "월 주거비와 이동 부담을 함께 비교"),
            ("후보 지역의 생활 인프라 확인이 번거로움", "병원·학교·공원 반경 집계", "생활권 접근성을 수치와 근거로 확인"),
            ("전세 계약 전 위험 요소 파악이 어려움", "전세가율·공시가격 대비 보증금·권리관계 체크리스트", "계약 전 확인 항목을 구조화"),
            ("추천 결과의 이유를 알기 어려움", "AI Agent 근거 그룹 답변", "가격, 통근, 위험, 입지 기준으로 설명"),
            ("후보를 여러 개 비교하기 어려움", "즐겨찾기 비교 패널", "단지별 가격·위험도·SOC 비교"),
        ],
    )

    add_heading(doc, "차별성", 2)
    add_paragraph(
        doc,
        "기존 부동산 플랫폼은 매물 검색과 가격 확인 중심이지만, MoveValue는 실제 거주 이후 발생하는 이동 부담, 생활 접근성, 주거 안정성까지 함께 분석합니다. 또한 추천 결과를 블랙박스로 보여주지 않고, 주거비·교통·생활 SOC·안전·전세 위험도 등 항목별 근거를 설명합니다.",
    )
    add_paragraph(
        doc,
        "생활권 추천과 전세 위험 점검을 결합해 사용자가 '어디에 살지'와 '계약해도 안전한지'를 함께 판단할 수 있도록 지원합니다.",
    )

    add_heading(doc, "데이터 출처 및 활용 목적 표", 2)
    add_table(
        doc,
        ["데이터/API", "활용 목적", "현재 상태"],
        [
            ("서울시 부동산 전월세가 정보 2025", "생활권별 월세·보증금·전세 중앙값, 예시 거래", "114,319건 집계 완료"),
            ("서울시 공동주택 아파트 정보 OpenAptInfo", "단지명, 주소, 좌표, 세대수, 사용승인일", "스냅샷 45건, 키 연결 시 전체 live"),
            ("국토교통부 아파트 매매 실거래가", "최근 매매가 live 보정", "어댑터 구현, 키 필요"),
            ("국토교통부 아파트 전월세 실거래가", "최근 전세가·월세 live 보정", "어댑터 구현, 키 필요"),
            ("공동주택 공시가격", "공시가격 대비 보증금 비율", "PNU/식별자 매핑 예정"),
            ("건축HUB 건축물대장", "층수, 용도, 구조, 위반건축물 확인", "확장 예정"),
            ("Kakao/ODsay/TMAP", "주소 검색과 대중교통 경로 검증", "어댑터 구현, 키 필요"),
            ("VWorld", "주소 좌표 보정과 용도지역 확장", "확장 예정"),
        ],
        widths=[2.0, 2.4, 2.2],
    )

    add_heading(doc, "서비스 기능별 구현 현황 표", 2)
    add_table(
        doc,
        ["기능", "구현 상태", "근거 파일"],
        [
            ("지도 기반 아파트 매칭", "구현 완료", "app/app.js, api/movevalue_api.py"),
            ("부동산 상세 대시보드", "구현 완료", "api/property_model.py, app/app.js"),
            ("전세 위험 신호", "구현 완료", "api/property_model.py"),
            ("AI Agent 질의응답", "구현 완료", "api/property_model.py, app/app.js"),
            ("후보 비교", "구현 완료", "app/app.js"),
            ("대중교통 live 경로", "키 기반 구조 구현", "api/route_adapters.py"),
            ("MOLIT 실거래가 보정", "키 기반 구조 구현", "api/real_estate_price_adapters.py"),
            ("공시가격/건축물대장", "확장 예정", ".env.example, docs/data-sources.md"),
        ],
    )

    add_heading(doc, "점수 산정 항목 및 가중치 표", 2)
    add_table(
        doc,
        ["항목", "기본 가중치", "산정 근거"],
        [
            ("통근", "25%", "목적지까지의 예상 대중교통 시간"),
            ("주거비", "25%", "사용자 예산 대비 월세/보증금 수준"),
            ("생활 SOC", "25%", "병원, 학교, 공원 접근성과 가구 유형별 선호"),
            ("안전", "25%", "치안시설, CCTV, 대기, 녹지, 전세 위험 신호 일부 반영"),
        ],
    )

    add_heading(doc, "전세 위험 신호 판단 기준표", 2)
    add_table(
        doc,
        ["위험 신호", "판단 기준", "표현 원칙"],
        [
            ("전세가율", "매매가 대비 전세가 비율이 높을수록 주의", "위험 신호 점검"),
            ("공시가격 대비 보증금", "공시가격 대비 보증금 비율이 높을수록 확인 필요", "계약 전 확인 필요"),
            ("주변 시세 대비 차이", "주변 전세 중앙값보다 과도하게 높은 경우 주의", "주의 요소"),
            ("최근 가격 변동성", "단기간 급등락 추정 시 재검증 필요", "실거래 API 연계 전 추정"),
            ("건축물 노후도", "노후 단지는 수선비·설비·주차 확인 필요", "확인 항목"),
            ("등기부 권리관계", "근저당, 압류, 선순위 임차권은 자동 판정하지 않음", "사용자 서류 확인"),
        ],
    )

    add_heading(doc, "시스템 구조와 API 연계", 2)
    add_paragraph(doc, "전체 Mermaid 다이어그램 9종은 `docs/mermaid-diagrams.md`에 정리했습니다. 핵심 구조는 아래와 같습니다.")
    add_table(
        doc,
        ["계층", "역할", "주요 API/파일"],
        [
            ("Web UI", "조건 입력, 지도, 상세 패널, 비교, Agent UI", "app/index.html, app/app.js, app/styles.css"),
            ("API Router", "정적 파일과 JSON API 제공", "api/movevalue_api.py"),
            ("Route Adapter", "Kakao/ODsay/TMAP/폴백 경로", "api/route_adapters.py"),
            ("Apartment Adapter", "OpenAptInfo live/스냅샷 단지 레이어", "api/apartment_adapters.py"),
            ("Price Adapter", "MOLIT 매매·전월세 live 보정", "api/real_estate_price_adapters.py"),
            ("Property Model", "가격, 거래 추이, 위험 신호, AI 요약, Agent 응답", "api/property_model.py"),
            ("Data Build", "전월세/SOC/안전 데이터 정규화", "scripts/build_real_dataset.py"),
        ],
    )

    add_heading(doc, "현황 및 구체화 방안", 2)
    add_paragraph(
        doc,
        "현재 프로토타입은 서울시 2025 전월세 실데이터 114,319건과 아파트 단지 스냅샷을 기반으로 동작합니다. 외부 API 키가 없을 때도 화면과 API가 동작하며, 응답에는 live/폴백/추정/확장 예정 상태가 함께 표시됩니다.",
    )
    add_paragraph(
        doc,
        "선정 후에는 ODsay/TMAP/Kakao 키 기반 live 검증, 서울 OpenAptInfo 전체 단지 로딩, 국토교통부 실거래가 단지명 매칭 검증, 공동주택 공시가격 PNU 매핑, 건축HUB 건축물대장 연계를 우선 추진합니다.",
    )

    add_heading(doc, "서비스 구현을 위한 협력 방안", 2)
    add_table(
        doc,
        ["협력 주체", "역할", "협력 가치"],
        [
            ("공공데이터 제공기관", "실거래가, 공시가격, 건축물대장, 교통 데이터 제공", "데이터 품질 개선과 활용사례 창출"),
            ("부동산 플랫폼", "사용자 접점과 매물/단지 정보 연계", "매물 검색 이상의 의사결정 보조 기능 제공"),
            ("MaaS/교통 데이터 사업자", "실제 경로, 교통비, 혼잡도 데이터 보강", "주거 선택과 이동 부담의 통합 분석"),
            ("지자체", "생활권 정책 실증과 사용자 피드백", "청년·신혼·고령자 생활권 정책 근거 확보"),
        ],
    )

    add_heading(doc, "기대효과 표", 2)
    add_table(
        doc,
        ["대상", "기대효과"],
        [
            ("시민", "주거비, 통근, 생활 인프라, 전세 위험 신호를 한 화면에서 비교해 선택 시행착오 감소"),
            ("지자체", "생활권별 주거·이동 부담과 SOC 격차를 정책 지표로 활용"),
            ("민간 플랫폼", "공공데이터 기반 추천·위험 점검 API로 서비스 차별화"),
            ("데이터 생태계", "국토교통 공공데이터를 실제 의사결정 서비스로 연결"),
        ],
    )

    add_heading(doc, "실증 추진 일정표", 2)
    add_table(
        doc,
        ["기간", "주요 작업", "산출물"],
        [
            ("2026.06", "API 키 기반 live 검증, 서울 전체 단지 로딩", "검증 로그, 데이터 갱신 배치"),
            ("2026.07", "실거래가/공시가격/건축물대장 정밀화, 위험 신호 고도화", "상세 대시보드 v2"),
            ("2026.08", "사용자 테스트, 정책/B2B 리포트 예시, 발표자료", "실증 리포트, 최종 프로토타입"),
        ],
    )

    add_heading(doc, "예산 사용 계획표", 2)
    add_table(
        doc,
        ["항목", "금액", "내용"],
        [
            ("데이터 수집·정제", "5,000천원", "실거래가·공시가격·건축물대장·SOC API 연계, 데이터 품질 점검"),
            ("프로토타입 개발", "8,000천원", "웹 UI, 점수 엔진, 통근·가격 API 어댑터, AI Agent 근거 구조화"),
            ("지도·분석 고도화", "3,000천원", "부동산 지도 라벨, 생활 SOC 반경, 단지 비교, 위험 신호 대시보드"),
            ("실증·사용자 테스트", "2,000천원", "시민 테스트, 인터뷰, 피드백 반영"),
            ("문서화·발표", "2,000천원", "기획서, 리포트, 발표자료"),
        ],
    )

    add_heading(doc, "확장 가능성 로드맵", 2)
    add_table(
        doc,
        ["단계", "확장 방향", "목표"],
        [
            ("1단계", "실제 API live 검증과 데이터 정밀화", "심사/실증 신뢰도 확보"),
            ("2단계", "지도 UX와 위험 신호 고도화", "실사용 가능한 주거 의사결정 화면"),
            ("3단계", "B2B/B2G API와 리포트", "민간 플랫폼과 지자체 연계"),
            ("4단계", "iOS SwiftUI 앱", "모바일 생활권 추천 서비스 확장"),
        ],
    )

    add_heading(doc, "3. 참가 서약서 확인", 1)
    add_paragraph(
        doc,
        "본 서비스는 타 대회 입상작 또는 타 서비스의 데이터를 무단 복제하지 않고, 공개 API/공공데이터와 직접 구현한 프로토타입을 기반으로 작성했습니다. 민간 부동산 서비스는 화면 구조와 사용자 경험만 참고하며, 이용약관·robots.txt·저작권·개인정보 이슈가 있는 크롤링은 수행하지 않습니다.",
    )
    add_kv_table(
        doc,
        [
            ("기관명", "이승보 / MoveValue 프로젝트팀"),
            ("신청자(대표) 성명", "이승보"),
            ("서명", "제출 전 직접 서명"),
        ],
    )

    add_heading(doc, "4. 개인정보 수집·이용 동의서 확인", 1)
    add_paragraph(
        doc,
        "공모전 신청자 개인정보 항목은 제공된 양식에 따라 신청자가 직접 확인하고 동의 여부를 표시해야 합니다. 본 보강본에는 실제 전화번호, 주소, 서명을 임의로 기재하지 않았습니다.",
    )
    add_kv_table(
        doc,
        [
            ("개인정보 수집·이용 동의", "제출 전 본인이 직접 확인 후 체크"),
            ("신청인 서명", "제출 전 직접 서명"),
            ("제출 전 확인", "연락처, 이메일, 서명일, 개인정보 동의 체크 여부"),
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
